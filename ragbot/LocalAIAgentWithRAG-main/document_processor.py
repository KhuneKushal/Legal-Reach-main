"""
Document Processor for Legal Reach RAG
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from datetime import datetime

# Document processing libraries
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Config
from config import DEFAULT_CONFIG, LEGAL_PDFS_DIR, LEGAL_DOCX_DIR, LEGAL_TXT_DIR, LOGS_DIR

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOGS_DIR / "document_processor.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Main class for processing legal documents
    Supports: PDF, DOCX, TXT
    """
    
    def __init__(self, config=None):
        """
        Initialize document processor
        
        Args:
            config: LegalReachConfig instance (uses DEFAULT_CONFIG if None)
        """
        self.config = config or DEFAULT_CONFIG
        self.supported_formats = self.config.document.supported_formats
        self.max_file_size = self.config.document.max_file_size
        logger.info(f" DocumentProcessor initialized with config: {config}")
    
    #  FILE LOADING 
    
    def load_documents_from_folder(self, folder_path: Path) -> List[Tuple[Document, str]]:
        """
        Load all documents from a folder
        
        Args:
            folder_path: Path to folder containing documents
            
        Returns:
            List of (Document, file_path) tuples
            
        Example:
            docs = processor.load_documents_from_folder(LEGAL_PDFS_DIR)
        """
        documents = []
        folder = Path(folder_path)
        
        if not folder.exists():
            logger.warning(f" Folder does not exist: {folder}")
            return documents
        
        # Find all supported file types
        for ext in self.supported_formats:
            for file_path in folder.glob(f"*{ext}"):
                try:
                    # Check file size
                    if file_path.stat().st_size > self.max_file_size:
                        logger.warning(f" File too large (>{self.max_file_size/1e6:.1f}MB): {file_path.name}")
                        continue
                    
                    # Load based on format
                    if ext.lower() == '.pdf':
                        docs = self._load_pdf(file_path)
                    elif ext.lower() == '.docx':
                        docs = self._load_docx(file_path)
                    elif ext.lower() == '.txt':
                        docs = self._load_txt(file_path)
                    else:
                        continue
                    
                    # Add to list with source path
                    for doc in docs:
                        documents.append((doc, str(file_path)))
                    
                    logger.info(f" Loaded: {file_path.name} ({len(docs)} pages/sections)")
                    
                except Exception as e:
                    logger.error(f"Failed to load {file_path.name}: {e}")
                    continue
        
        logger.info(f" Total documents loaded: {len(documents)}")
        return documents
    
    def _load_pdf(self, file_path: Path) -> List[Document]:
        """
        Load PDF file using pypdf
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of Document objects (one per page)
        """
        documents = []
        
        try:
            pdf_reader = PdfReader(file_path)
            num_pages = len(pdf_reader.pages)
            
            logger.debug(f"📄 PDF has {num_pages} pages: {file_path.name}")
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    text = page.extract_text()
                    
                    if text.strip():  # Only add if text exists
                        doc = Document(
                            page_content=text,
                            metadata={
                                "source": file_path.name,
                                "page": page_num,
                                "total_pages": num_pages,
                                "format": "pdf",
                                "upload_date": datetime.now().isoformat()
                            }
                        )
                        documents.append(doc)
                except Exception as e:
                    logger.warning(f"⚠️ Failed to extract text from page {page_num}: {e}")
                    continue
        
        except Exception as e:
            logger.error(f"❌ Error reading PDF {file_path.name}: {e}")
        
        return documents
    
    def _load_docx(self, file_path: Path) -> List[Document]:
        """
        Load DOCX file using python-docx
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of Document objects
        """
        documents = []
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract all paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Combine into sections
            full_text = "\n".join(text_parts)
            
            if full_text.strip():
                document = Document(
                    page_content=full_text,
                    metadata={
                        "source": file_path.name,
                        "format": "docx",
                        "upload_date": datetime.now().isoformat(),
                        "total_pages": len(doc.paragraphs)
                    }
                )
                documents.append(document)
            
            logger.debug(f" DOCX loaded: {file_path.name}")
        
        except Exception as e:
            logger.error(f" Error reading DOCX {file_path.name}: {e}")
        
        return documents
    
    def _load_txt(self, file_path: Path) -> List[Document]:
        """
        Load TXT file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            List of Document objects
        """
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if text.strip():
                document = Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "format": "txt",
                        "upload_date": datetime.now().isoformat()
                    }
                )
                documents.append(document)
            
            logger.debug(f" TXT loaded: {file_path.name}")
        
        except Exception as e:
            logger.error(f" Error reading TXT {file_path.name}: {e}")
        
        return documents
    
    # ==================== CHUNKING ====================
    
    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        """
        Split documents into chunks with overlap
        Preserves metadata and source information
        
        Args:
            documents: List of Document objects
            
        Returns:
            List of chunked Document objects
            
        Example:
            chunks = processor.chunk_documents(documents)
        """
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.document.chunk_size,
            chunk_overlap=self.config.document.chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]  # Try to split intelligently
        )
        
        chunked_docs = []
        total_chunks = 0
        
        for doc in documents:
            try:
                # Split the document
                chunks = splitter.split_text(doc.page_content)
                
                # Create Document objects for each chunk with metadata
                for chunk_idx, chunk_text in enumerate(chunks, 1):
                    chunk_doc = Document(
                        page_content=chunk_text,
                        metadata={
                            **doc.metadata,  # Keep original metadata
                            "chunk": chunk_idx,
                            "total_chunks": len(chunks)
                        }
                    )
                    chunked_docs.append(chunk_doc)
                
                total_chunks += len(chunks)
                logger.debug(f"  {doc.metadata.get('source', 'unknown')} → {len(chunks)} chunks")
            
            except Exception as e:
                logger.error(f" Error chunking document: {e}")
                continue
        
        logger.info(f"Total chunks created: {total_chunks} from {len(documents)} documents")
        return chunked_docs
    
    # ----------------------------PREPROCESSING ---------------------------
    
    def preprocess_text(self, text: str) -> str:
        """
        Clean and preprocess document text
        
        Args:
            text: Raw text from document
            
        Returns:
            Cleaned text
        """
        # Remove extra whitespace
        text = " ".join(text.split())
        
        # Remove common artifacts
        text = text.replace("Page 1", "").replace("© ", "")
        
        # Normalize quotes
        text = text.replace(""", '"').replace(""", '"')
        text = text.replace("'", "'")
        
        return text.strip()
    
    # -------------------STATISTICS 
    
    def get_document_stats(self, documents: List[Document]) -> Dict:
        """
        Get statistics about loaded documents
        
        Args:
            documents: List of Document objects
            
        Returns:
            Dictionary with statistics
        """
        total_chars = sum(len(doc.page_content) for doc in documents)
        total_tokens_estimate = total_chars / 4  # Rough estimate: 1 token ≈ 4 chars
        
        sources = set(doc.metadata.get("source", "unknown") for doc in documents)
        formats = set(doc.metadata.get("format", "unknown") for doc in documents)
        
        stats = {
            "total_documents": len(documents),
            "total_characters": total_chars,
            "estimated_tokens": total_tokens_estimate,
            "unique_sources": list(sources),
            "formats": list(formats),
            "avg_chars_per_doc": total_chars / len(documents) if documents else 0
        }
        
        logger.info(f"📊 Document Statistics: {stats}")
        return stats


# ----------------------------CONVENIENCE FUNCTIONS 

def load_all_legal_documents() -> List[Document]:
    """
    Load all legal documents from the legal_documents folder
    
    Returns:
        List of Document objects (not chunked)
    """
    processor = DocumentProcessor()
    
    documents = []
    
    # Load from all format directories
    for folder in [LEGAL_PDFS_DIR, LEGAL_DOCX_DIR, LEGAL_TXT_DIR]:
        if folder.exists():
            docs_with_paths = processor.load_documents_from_folder(folder)
            documents.extend([doc for doc, _ in docs_with_paths])
    
    logger.info(f"✅ Loaded {len(documents)} total documents")
    return documents


def process_legal_documents() -> List[Document]:
    """
    Load and chunk all legal documents
    This is the main entry point for processing
    
    Returns:
        List of chunked Document objects ready for embedding
    """
    processor = DocumentProcessor()
    
    # Load all documents
    logger.info("📂 Loading legal documents...")
    raw_documents = load_all_legal_documents()
    
    if not raw_documents:
        logger.warning("⚠️ No documents found!")
        return []
    
    # Chunk documents
    logger.info("  Chunking documents...")
    chunked_documents = processor.chunk_documents(raw_documents)
    
    # Show stats
    stats = processor.get_document_stats(chunked_documents)
    
    return chunked_documents


#---------------------TESTING ---------------------------

if __name__ == "__main__":
    """Test document processor"""
    logger.info(" Testing DocumentProcessor...")
    
    # Create sample test file
    test_text = """
    LEGAL AGREEMENT
    
    This Agreement is entered into between Party A and Party B.
    
    1. DEFINITIONS
    For purposes of this Agreement, the following terms shall have the meanings set forth below.
    
    2. OBLIGATIONS
    Party A shall provide services as outlined in Schedule A.
    Party B shall make payments as outlined in Schedule B.
    
    3. TERM
    This Agreement shall commence on the effective date and continue for one (1) year.
    """
    
    # Save test file
    test_path = LEGAL_TXT_DIR / "test_legal_document.txt"
    test_path.write_text(test_text)
    logger.info(f"📝 Created test document: {test_path}")
    
    # Process it
    processor = DocumentProcessor()
    raw_docs = processor.load_documents_from_folder(LEGAL_TXT_DIR)
    logger.info(f"📄 Raw documents: {len(raw_docs)}")
    
    docs_list = [doc for doc, _ in raw_docs]
    chunked = processor.chunk_documents(docs_list)
    logger.info(f"✂️  Chunked documents: {len(chunked)}")
    
    for idx, chunk in enumerate(chunked[:2], 1):
        logger.info(f"\nChunk {idx}:")
        logger.info(f"  Source: {chunk.metadata.get('source')}")
        logger.info(f"  Text: {chunk.page_content[:100]}...")
    
    logger.info("✅ DocumentProcessor test complete!")
